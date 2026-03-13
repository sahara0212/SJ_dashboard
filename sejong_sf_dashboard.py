import streamlit as st
import pandas as pd
import plotly.express as px
import time
import graphviz
from docx import Document
from io import BytesIO

# -----------------------------------------------------------------------------
# 1. 페이지 및 스타일 설정
# -----------------------------------------------------------------------------
st.set_page_config(page_title="SHIN & KIM | SF2 Dashboard", layout="wide")

# 사이드바: 프로세스 흐름형 메뉴 구성
st.sidebar.image("https://www.shinkim.com/assets/images/common/logo_ko.png", width=200)
st.sidebar.header("SF2 Service Flow")
menu = st.sidebar.radio("진행 단계", [
    "1. 홈 (Home)", 
    "2. [Step 1] 자가진단 (Excel)", 
    "3. [Step 2] 아키텍처 (PPT)", 
    "4. [Step 3] 신청서 생성 (Word)"
])

st.sidebar.info("""
**💡 SF2 패키지 구성품**
1. 자가진단 키트 (.xlsx)
2. 표준 보안 아키텍처 (.pptx)
3. 혁신금융 신청서 템플릿 (.docx)
""")

# -----------------------------------------------------------------------------
# 함수: 파일 생성 유틸리티
# -----------------------------------------------------------------------------
# 진짜 워드 파일 생성
def generate_word_file(company_name, service_name, protection_text):
    doc = Document()
    doc.add_heading('혁신금융서비스 지정 신청서', 0)
    doc.add_heading('1. 신청인 일반 현황', level=1)
    doc.add_paragraph(f'• 회사명: {company_name}')
    doc.add_paragraph(f'• 서비스 명칭: {service_name}')
    doc.add_heading('2. 서비스 주요 내용', level=1)
    doc.add_paragraph('본 서비스는 생성형 AI(MS Copilot)를 활용하여 금융 업무 효율성을 극대화합니다.')
    doc.add_heading('3. 소비자 보호 및 보안 대책', level=1)
    doc.add_paragraph(protection_text)
    doc.add_paragraph(f'\n2025년 X월 X일\n신청인: {company_name} 대표이사 (인)')
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# (데모용) 엑셀/PPT 다운로드 시뮬레이션
def get_dummy_file(content_type):
    return BytesIO(b"Demo File Content")

# -----------------------------------------------------------------------------
# 2. 메인 홈 (Home)
# -----------------------------------------------------------------------------
if menu == "1. 홈 (Home)":
    st.title("🛡️ SaaS Fast-Track for Finance (SF2)")
    st.markdown("### 규제 샌드박스 신청부터 보안 아키텍처 수립까지, 금융권 특화 턴키 솔루션")
    
    col1, col2 = st.columns([2, 1])
    with col1:
        st.markdown("""
        **법무법인 세종이 제공하는 SF2 패키지는 다음 3가지를 한 번에 해결합니다.**
        
        1. **자가진단 (Self-Check):** 복잡한 데이터 리스크를 엑셀 하나로 자동 분석
        2. **보안 설계 (Architecture):** 금융보안원 심의를 통과하는 표준 아키텍처 제공
        3. **신청 대행 (Application):** 수주율 높은 혁신금융서비스 신청서 자동 생성
        """)
        
    with col2:
        st.success("✅ **지금 바로 프로젝트를 시작하세요.**")
        if st.button("🚀 SF2 프로세스 시작하기"):
            with st.spinner("솔루션 모듈 로딩 중..."):
                time.sleep(1)
            st.info("좌측 메뉴의 **'[Step 1] 자가진단'**으로 이동해주세요.")

# -----------------------------------------------------------------------------
# 3. [Step 1] 자가진단 (Excel) - 리스크 분석
# -----------------------------------------------------------------------------
elif menu == "2. [Step 1] 자가진단 (Excel)":
    st.title("📊 Step 1. 데이터 라이프사이클 자가진단")
    st.markdown("제공된 엑셀 템플릿을 작성하여 업로드하면, **잠재된 법적 리스크를 자동 분석**합니다.")

    # 1. 템플릿 다운로드 섹션
    with st.expander("📥 1. 자가진단 키트(Excel) 다운로드", expanded=True):
        st.write("금융권 데이터 컴플라이언스 주요 점검 항목 50개가 포함된 엑셀 파일입니다.")
        st.download_button("엑셀 템플릿 다운로드 (.xlsx)", data=get_dummy_file("excel"), file_name="SF2_자가진단키트_v1.0.xlsx")

    # 2. 업로드 및 분석 섹션
    st.markdown("---")
    st.markdown("#### 📤 2. 진단 결과 업로드 및 분석")
    uploaded_file = st.file_uploader("작성된 엑셀 파일을 업로드하세요.", type=['xlsx'])

    # (데모용) 파일을 올리거나 버튼을 누르면 분석 결과 표시
    if uploaded_file is not None or st.button("데모 데이터로 분석 실행"):
        with st.spinner("데이터 정합성 및 파기 의무 준수 여부 분석 중..."):
            time.sleep(1.5)
        
        st.error("🚨 **진단 결과: 고위험 (High Risk)** - 즉각적인 조치가 필요합니다.")
        
        # KPI 지표 (KB라이프 데이터 활용)
        c1, c2, c3 = st.columns(3)
        c1.metric("미파기(좀비) 테이블", "288 개", "-23% (위험)", delta_color="inverse")
        c2.metric("데이터 정합성 오류", "596 건", "-1.3% (주의)", delta_color="inverse")
        c3.metric("데이터 흐름 복잡도", "83 종", "관리 필요", delta_color="off")
        
        # 차트
        c1, c2 = st.columns(2)
        with c1:
            df = pd.DataFrame({'상태': ['파기 완료', '미파기'], '개수': [958, 288]})
            fig = px.pie(df, values='개수', names='상태', title='개인정보 파기 의무 준수율', color_discrete_sequence=['lightgrey', 'red'])
            st.plotly_chart(fig, use_container_width=True)
        
        with c2:
            st.info("""
            **💡 세종의 솔루션:**
            - **미파기 데이터:** 별도 분리 보관(Cold Storage) 및 접근 통제 적용
            - **오류 데이터:** 마케팅 활용 정지(Freeze) 처리로 법적 리스크 헷지
            """)

# -----------------------------------------------------------------------------
# 4. [Step 2] 보안 아키텍처 (PPT) - 설계
# -----------------------------------------------------------------------------
elif menu == "3. [Step 2] 아키텍처 (PPT)":
    st.title("🛡️ Step 2. 표준 보안 아키텍처 수립")
    st.markdown("진단된 리스크를 통제하고 **망분리 규제 예외를 인정받기 위한 표준 설계도**입니다.")

    col1, col2 = st.columns([3, 1])
    
    with col1:
        st.markdown("##### 📌 SF2 논리적 망분리 구성도 (MS Copilot 연계)")
        try:
            graph = graphviz.Digraph()
            graph.attr(rankdir='LR')
            graph.node('User', '업무 단말(VDI)', shape='box', style='filled', fillcolor='lightblue')
            graph.node('Internal', '내부망', shape='box')
            graph.node('Gateway', 'SF2 보안 게이트웨이\n[PII 마스킹/통제]', shape='doubleoctagon', style='filled', fillcolor='orange', fontcolor='black')
            graph.node('SaaS', 'MS 365 Cloud', shape='cloud', style='filled', fillcolor='lightgrey')
            
            graph.edge('User', 'Internal')
            graph.edge('Internal', 'Gateway', label='전송')
            graph.edge('Gateway', 'SaaS', label='안전 전송', color='blue')
            graph.edge('Gateway', 'Internal', label='차단(민감정보)', color='red', style='dashed')
            st.graphviz_chart(graph)
        except:
            st.warning("(Graphviz가 설치되지 않아 다이어그램 대신 텍스트로 표시합니다.)\n[단말기] -> [내부망] -> [보안 게이트웨이] -> [MS 클라우드]")

    with col2:
        st.markdown("#### 📥 산출물 다운로드")
        st.write("금융보안원 심의 제출용 상세 설계도(PPT)를 다운로드하세요.")
        st.download_button("표준 아키텍처.pptx", data=get_dummy_file("ppt"), file_name="SF2_표준아키텍처_v1.0.pptx")
        
        st.success("""
        **핵심 보안 요소:**
        1. **Gateway:** 주민번호 자동 마스킹
        2. **Log:** 전송 데이터 감사 기록
        3. **Control:** VDI 환경 접근 통제
        """)

# -----------------------------------------------------------------------------
# 5. [Step 3] 신청서 생성 (Word) - 최종 산출물
# -----------------------------------------------------------------------------
elif menu == "4. [Step 3] 신청서 생성 (Word)":
    st.title("📝 Step 3. 혁신금융서비스 신청서 생성")
    st.markdown("앞 단계의 진단 및 설계 결과를 반영하여 **금융위원회 제출용 신청서를 자동 완성**합니다.")

    with st.form("app_form"):
        c1, c2 = st.columns(2)
        company = c1.text_input("신청 회사명", value="KB라이프")
        service = c2.text_input("서비스 명칭", value="생성형 AI 기반 업무 비서")
        
        st.markdown("**자동 맵핑된 소비자 보호 및 보안 대책**")
        # Step 1, 2의 결과가 자동으로 반영된 것처럼 텍스트 생성
        protection_plan = st.text_area("보호 방안 (수정 가능)", 
            value="1. [Step 2] SF2 보안 게이트웨이를 통해 개인식별정보(주민번호 등)를 실시간으로 탐지하여 자동 비식별화(마스킹) 처리함.\n"
                  "2. [Step 1] 진단 결과 식별된 288개의 미파기 테이블은 별도 망으로 분리하여 LLM 학습 접근을 원천 차단함.\n"
                  "3. 데이터 정합성 오류(1.3%) 데이터는 마케팅 활용 대상에서 즉시 제외하여 오발송 리스크를 방지함.", 
            height=150)
            
        submit = st.form_submit_button("📄 신청서 최종 생성")

    if submit:
        with st.spinner("법률 검토 및 문서 생성 중..."):
            time.sleep(2)
        
        # 진짜 워드 파일 생성
        word_data = generate_word_file(company, service, protection_plan)
        
        st.success("문서 생성이 완료되었습니다! 아래 버튼을 눌러 확인하세요.")
        st.download_button(
            label="📥 혁신금융서비스 신청서(.docx) 다운로드",
            data=word_data,
            file_name=f"{company}_혁신금융신청서.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
